import math
from scipy.fftpack import fft,  ifft
import numpy as np
import matplotlib.pyplot as plt
import scipy.special as ss
import docx
import numpy.matlib
import matplotlib.lines as mlines

def func(x):  #функция апроксимируемая рядом фурье
    x= x%1
    if (abs(x)<=1/2):
        return 1
    return -1

#--------------------------------------------------------------------------------------------
# Вспомогательные функции для расчета ряда Фурье
def cn(n):
   c = ym*np.exp(-1j*2*n*np.pi*xm/period)
   return c.sum()/c.size

def f_okon(x, h, name):
    f = np.array([name(i,h)*2*cn(i)*np.exp(1j*2*i*np.pi*x/period) for i in range(0,h)])
    return f.sum()

def f_okon_l(x, h, name, l):
    f = np.array([name(i, h, l)*2*cn(i)*np.exp(1j*2*i*np.pi*x/period) for i in range(0,h)])
    return f.sum()

#-----------------------------------------------------------------------------------------
# Оконные функции
# i - Аргумент оконной функции
# h - Размер окна
# b - дополнительный параметр
def priamougl_1(i , h):  #Прямоугольное окно
    if ((0<=i) and (i<h)):
        return 1
    return 0 

def Bartlett_1(i , h): # Окно Бартлетта 
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return((h - 2*abs(j-h/2))/h)
    return 0 

def Hann_1(i, h): # Окно Ханна
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return(0.5 - 0.5*math.cos(2*math.pi*j/h))
    return 0 

def Hann(i, h): # Окно Ханна
    if ((0<=i) and (i<h)):
        return(0.5 - 0.5*math.cos(2*math.pi*i/h))
    return 0 

def Hamming_1(i, h): # Окно Хемминга
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return(0.54 - 0.46*math.cos(2*math.pi*j/h))
    return 0 

def Bartlett_Hann_1(i, h): # Окно Бартлетта-Ханна
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return(0.62 - 0.48*abs(j/h-1/2) - 0.38*math.cos(2*math.pi*j/h))
    return 0

def blackman_haris_1(i,h): # Окно Блэкмана-Натталла
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return (0.35875 - 0.48829*math.cos(2*math.pi*j/h) + 0.14128*math.cos(4*math.pi*j/h) - 0.01168*math.cos(6*math.pi*j/h))
    return 0 

def Nattal_1(i,h): #Окно Натталла
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return (0.355768 - 0.487396*math.cos(2*math.pi*j/h) + 0.144232*math.cos(4*math.pi*j/h) - 0.012604*math.cos(6*math.pi*j/h))
    return 0 

def blackman_nattal_1(i,h): # Окно Блэкмана-Натталла
    j = i/2 + h/2
    if ((0<=i) and (i<h)):
        return (0.3635819 - 0.4891775*math.cos(2*math.pi*j/h) + 0.1365995*math.cos(4*math.pi*j/h) - 0.0106411*math.cos(6*math.pi*j/h))
    return 0 

def Kayser_1(i, h, b): # Окно Кайзера
    j = i/2 + h/2
    #j = i
    if ((0<=i) and (i<h)):
        z = b*(1-(2*(j - h/2)/h)**2)**(1/2)
        return (ss.i0(z)/ss.i0(b))
    return 0 

def Ghaus_1(i, h, b): # Окно Гаусса
    j = i/2 + h/2
    #j = i
    if ((0<=i) and (i<h)):
        z = np.exp((-2*(j-h/2)**2)/(b*h)**2)
        return z
    return 0 

#------------------------------------------------------------------------------------

period=1
xm = np.arange(-2, 2+0.001, 0.001) #Вспомогательная часть для расчетов 
Nh = int(4/0.001)
ym = np.array([func(x).real for x in xm])

def analys(x, y): #Расчет площади разности изначального и сглаженного графика
    i = 0
    j = 0
    ny = np.array([])
    nx = np.array([])
    hx = x[1] - x[0]
    while ((y[i]<1) and (x[i] <= 0.5)):
        i+=1
    
    while (x[j]<0.5):
            j+=1
    j+=1

    y1 =  np.array([func(xi).real for xi in x[0:j]])

    if (x[i] < 0.5):
        je = j
        while (y[j]<1):
            j-=1
        j+=1
        Gp_S = np.trapz(np.absolute(y[i:j+1] - 1), dx = hx)
        Sd_S = np.trapz(1 - y[0:i], dx = hx) + np.trapz(1 - y[j:je], dx = hx)
        """
        plt.plot(x[0:je], y1, color='b') 
        plt.plot(x[0:je], y[0:je], color='r')
        plt.fill_between(x[0:i], y[0:i], y1[0:1], facecolor='green')
        plt.fill_between(x[j:je], y[j:je], y1[j:je], facecolor='pink')
        plt.fill_between(x[i:j], y[i:j], y1[i:j], facecolor='yellow')
        plt.grid()
        plt.show()
        """
    else:
        Gp_S = 0
        Sd_S = np.trapz(1-y[0:j], dx = hx)
        """
        plt.plot(x[0:j], y1, color='b') 
        plt.plot(x[0:j], y[0:j], color='r') 
        plt.fill_between(x[0:j], y[0:j], y1, facecolor='pink')
        plt.grid()
        plt.show() 
        """
    
    res = np.array([])
    res = np.append(res, Gp_S.real)
    res = np.append(res, Sd_S.real)
    res = np.append(res, Gp_S.real + Sd_S.real)
    return res

def analys_1(a, b, c, h): #Расчет площади разности изначального и сглаженного графика для без параметрических окон
    xm = np.arange(a, b+c, c)
    y1 = np.array([-f_okon(x, h, priamougl_1).real for x in xm])
    y2 = np.array([-f_okon(x, h, Bartlett_1).real for x in xm])
    y3 = np.array([-f_okon(x, h, Hann_1).real for x in xm])
    y4 = np.array([-f_okon(x, h, Hamming_1).real for x in xm])
    y5 = np.array([-f_okon(x, h, Bartlett_Hann_1).real for x in xm])
    y6 = np.array([-f_okon(x, h, blackman_haris_1).real for x in xm])
    y7 = np.array([-f_okon(x, h, Nattal_1).real for x in xm])
    y8 = np.array([-f_okon(x, h, blackman_nattal_1).real for x in xm])
    res1 = analys(xm, y1)
    res2 = analys(xm, y2)
    res3 = analys(xm, y3)
    res4 = analys(xm, y4)
    res5 = analys(xm, y5)
    res6 = analys(xm, y6)
    res7 = analys(xm, y7)
    res8 = analys(xm, y8)
    resall = [res1, res2, res3, res4, res5, res6, res7, res8]
    doc = docx.Document()
    table = doc.add_table(rows = 8, cols = 3)  
    table.style = 'Table Grid'
    for rows in range(0, 8):
        for col in range(0, 3):
            cell = table.cell(rows, col)
            cell.text = str(resall[rows][col]) 
    doc.save('S_nnon_parametr.docx')
    return None

def analys29(N, l):
    x2 = np.arange(0, 2**18, 1)
    x1 = np.linspace(0, 1, 2**18, endpoint=False)
    y1 = np.array([func(x) for x in x1])
    y1 = np.fft.fft(y1)
    res = np.array([])
    y = np.array([-Kayser_1(hi, 2**N, l)*y1[hi] for hi in x2])
    sp = np.fft.fft(y)
    sp = 2*sp/(2**18)
    res = analys(x1, sp)
    print(res[0], ' | ', res[1], ' | ', res[2], ' | ', N, ' | ', l)
    return res

def analys_graf_l(h1, h2, hc, l1, l2, lc): #Построение трехмерного графика для функции сглаженной параметрическм окном 
    l = np.arange(l2, l1, lc)
    h = np.arange(h2, h1, hc)
    G = np.zeros((len(h), len(l)))
    S = np.zeros((len(h), len(l)))
    sum = np.zeros((len(h), len(l)))
    xgrid, ygrid = np.meshgrid(l, h)
    fig = plt.figure()
    ax = fig.add_subplot(111, projection='3d')
    for hi in range(len(h)):
        for i in range(len(l)):
            res = analys29(h[hi], l[i])
            #G[hi][i] = res[0]
            #S[hi][i] = res[1]
            sum[hi][i] = 10*np.log10(res[2])
    #ax.plot_surface(ygrid, xgrid, G)
    #ax.plot_surface(ygrid, xgrid, S)
    ax.plot_surface(ygrid, xgrid, sum)
    ax.legend()
    plt.show()     
    return None

def analys_l1_li(x, y): #Расчет L_1 и L_inf разности изначального и сглаженного графика
    i = 0
    j = 0
    while ((y[i+1] - y[i])>0):
        i+=1
    while (x[j]<0.5):
        j+=1
    je = j
    while ((y[j+1] - y[j])<0):
        j-=1
    y6 = np.array([func(xi).real for xi in x[i:j+2]])
    y7 = np.array([func(xi).real for xi in x[0:je]])
    res = np.array([])
    res = np.append(res, (x[i]))
    res = np.append(res, (0.5 - x[j+1]))
    res = np.append(res, x[j+1] - x[i])
    res = np.append(res, np.linalg.norm((y[i:j+2].real - y6), 1))
    res = np.append(res, np.linalg.norm((y[i:j+2].real - y6), np.inf))
    res = np.append(res, np.linalg.norm((y[0:je].real - y7), 1))
    return res

def mod(a , b, c, h, name): #Построение изначального и сглаженного графика
    xm = np.arange(a, b+c, c)
    y1 = np.array([f_okon(x, h, name).real for x in xm])
    y2 = np.array([func(x).real for x in xm])
    plt.plot(xm, y1, 'r-', linewidth = 4) 
    plt.plot(xm, y2, 'b-', linewidth = 4, alpha = 0.6) 
    plt.grid()
    plt.show() 
    return None

def mod_l(a, b, c, N, name, l): # Функция которая показывает разницу между изначальной и сглаженной окном с параметром
    xm = np.arange(a, b+c, c)
    y2 = np.array([f_okon_l(x, N, name, l).real for x in xm])
    y3 = np.array([func(x).real for x in xm]) 
    z = abs(y2 - y3)
    print(name)
    print(np.linalg.norm(z, 1))
    print(np.linalg.norm(z, 2))
    print(np.linalg.norm(z, np.inf))    
    plt.plot(xm, y3)
    plt.plot(xm, y2) 
    plt.grid()
    plt.show()  
    return None 

def mod_3h(a, b, c, N1, N2, N3, name):
    xm = np.arange(a, b+c, c)
    y2 = np.array([f_okon_l(x, N1, name, 15).real for x in xm])
    y3 = np.array([f_okon_l(x, N2, name, 15).real for x in xm])
    y4 = np.array([f_okon_l(x, N3, name, 15).real for x in xm])
    y5 = np.array([func(x).real for x in xm])    
    plt.plot(xm, y2, color='r')
    plt.plot(xm, y3, color='g' )
    plt.plot(xm, y4, color='b') 
    plt.plot(xm, y5) 
    plt.grid()
    plt.show()  
    doc = docx.Document()
    table = doc.add_table(rows = 4, cols = 3)  
    table.style = 'Table Grid'   
    s = ['h','L_1', 'Среднеквадратическое отклонение']
    res = [N3, np.linalg.norm(abs(y4 - y5), 1), (((abs(y4 - y5))**2).sum()/len(y5))**(1/2), 
           N2, np.linalg.norm(abs(y3 - y5), 1), (((abs(y3 - y5))**2).sum()/len(y3))**(1/2), 
           N1, np.linalg.norm(abs(y2 - y5), 1), (((abs(y2 - y5))**2).sum()/len(y2))**(1/2)]
    for col in range(3):
        cell = table.cell(0, col)
        cell.text = str(s[col])
    for rows in range(1, 4):
        for col in range(3):
            cell = table.cell(rows, col)
            cell.text = str(res[(rows-1)*3 + col]) 
    doc.save('tablefunc.docx') 
    return None   

def analys_old(x, y): # Анализ оконных функций через нормы и размер зоны перехода
    i = 0
    j = 0
    ny = np.array([])
    nx = np.array([])
    while ((y[i+1] - y[i])>0):
        i+=1
    while (x[j]<0.5):
        j+=1
    while ((y[j+1] - y[j])<0):
        j-=1
    for k in range(i, j+2):
        ny = np.append(ny, y[k])
        nx = np.append(nx, x[k])
    y5 = np.array([func(xi).real for xi in x])
    y6 = np.array([func(xi).real for xi in nx])
    #plt.plot(x, y5, color='b') 
    #plt.plot(x, y, color='g', marker='o' )
    #plt.plot(nx, ny, color='r', marker='o')
    #plt.grid()
    #plt.show() 
    res = np.array([])
    res = np.append(res, (x[i]))
    res = np.append(res, (0.5 - x[j+1]))
    res = np.append(res, x[j+1] - x[i])
    res = np.append(res, np.linalg.norm(ny - y6 , 1))
    res = np.append(res, ((((ny - y6)**2).sum())/len(ny))**(1/2))
    res = np.append(res, np.linalg.norm(ny - y6, np.inf)) 
    return res

def all_analys(a , b, c, N1): # Анализ для нескольких оконных функций и создание ворд файла
    xm = np.arange(a, b+c, c)
    y1 = np.array([f_okon(x, N1, priamougl_1).real for x in xm])
    y2 = np.array([f_okon(x, N1, Bartlett_1).real for x in xm])
    y3 = np.array([f_okon(x, N1, Hann_1).real for x in xm])
    y4 = np.array([f_okon(x, N1, Hamming_1).real for x in xm])
    y5 = np.array([f_okon(x, N1, Bartlett_Hann_1).real for x in xm])
    y6 = np.array([f_okon(x, N1, blackman_haris_1).real for x in xm])
    y7 = np.array([f_okon(x, N1, blackman_nattal_1).real for x in xm])
    y8 = np.array([f_okon(x, N1, Nattal_1).real for x in xm])
    y9 = np.array([f_okon_l(x, N1, Ghaus_1, 0.5).real for x in xm])
    y10 = np.array([f_okon_l(x, N1, Ghaus_1, 0.22).real for x in xm])
    y11 = np.array([f_okon_l(x, N1, Kayser_1, 4).real for x in xm])
    y12 = np.array([f_okon_l(x, N1, Kayser_1, 12).real for x in xm])
    res1 = analys_old(xm, y1)
    print('1')
    res2 = analys_old(xm, y2)
    print('2')
    res3 = analys_old(xm, y3)
    print('3')
    res4 = analys_old(xm, y4)
    print('4')
    res5 = analys_old(xm, y5)
    print('5')
    res6 = analys_old(xm, y6)
    print('6')
    res7 = analys_old(xm, y7)
    print('7')
    res8 = analys_old(xm, y8)
    print('8')
    res9 = analys_old(xm, y9)
    print('9')
    res10 = analys_old(xm, y10)
    print('10')
    res11 = analys_old(xm, y11)
    print('11')
    res12 = analys_old(xm, y12)
    print('12')
    
    names = ['priam', 'Bart', 'Hann', 'Hamm', 
            'Bart-Hann', 'Black-Harr', 'Nattal', 'Black-Natt',
            'Ghaus(0.5)', 'Ghaus(0.22)', 'Kayser(4)', 'Kayser(12)']
    allres = np.array([res1, res2, res3, res4, res5, res6, res7, res8, res9, res10, res11, res12])

    doc = docx.Document()
    table = doc.add_table(rows = 13, cols = 7)  
    table.style = 'Table Grid'
    s = ['Окно', 'Левая переходная зона', 'Правая переходная зона', 'Зона осцилляций', 'L_1_oz', 'средне-квадратическое отклонение', 'абсолютное отклонение']
    for col in range(7):
        cell = table.cell(0, col)
        cell.text = str(s[col])
    for rows in range(1, 13):
        cell = table.cell(rows, 0)
        cell.text = str(names[rows-1])
        for col in range(1, 7):
            cell = table.cell(rows, col)
            cell.text = str(allres[rows-1][col-1]) 
    doc.save('alltablefunc.docx') 
    return None

def spect_plot(name, w, h): # вспомогательная функция для спектральной плотности энергиии для не параметрического окна
    z = 0
    z1 = 0
    for n in range(-h, h):
        z += name(n,h)*np.exp(-1j*w*n)
    for n in range(-h, h):
        z1 += name(n,h)*np.exp(-1j*0*n)
    z = z**2
    z1 = z1**2
    z1 = abs(z1)
    z = abs(z)
    z = z/z1
    z = 10*math.log10(z)
    return z

def spect_plot_fft(name, a, b, h, N): # Cпектральная плотность энергиии через бпф для не параметрического окна
    x = np.arange(h)
    x1 = np.linspace(a, b, h)
    y = np.array([name(hi, N) for hi in x])
    sp = np.fft.fft(y)
    sp = sp/sp[0]
    sp = abs(sp)
    sp = sp**2
    sp = 10*np.log10(sp)
    freq = np.fft.fftfreq(len(x), x[1] - x[0])
    plt.plot(freq, sp)
    plt.show()
    return None

def spect_plot3(name, w,  h, l): # вспомогательная функция для спектральной плотности энергиии для параметрического окна
    z = 0
    z1 = 0
    for n in range(-h, h):
        z += name(n,h,l)*np.exp(-1j*w*n)
    for n in range(-h, h):
        z1 += name(n,h,l)*np.exp(-1j*0*n)
    z1 = abs(z1)
    z = abs(z)
    z = z**2
    z1 = abs(z1)
    z1 = z1**2
    z = z/z1
    z = 10*math.log10(z)
    return z

def graf_mainSp(a, b, c, h, name): # Cпектральная плотность энергиии для не параметрического окна
    xm = np.arange(a, b+c, c)
    ym1 = np.array([spect_plot(name, x, h) for x in xm])
    print(ym1)
    plt.plot(xm, ym1 , color='b')
    plt.grid()
    plt.show() 
    return None

def graf_mainSp3(a, b, c, h, name, l): # Cпектральная плотность энергиии для параметрического окна
    xm = np.arange(a, b+c, c)
    ym1 = np.array([spect_plot3(name, x, h , l) for x in xm])
    plt.plot(xm, ym1 , color='b')
    plt.grid()
    plt.show() 
    return None
#-------------------------------------------------------------------------------------------------------------
#Сигма факторы
def sig(i, h):
    if (i==0):
        return 1
    t = (np.sin(np.pi*i/h)/(np.pi*i/h))
    return t

def f_okon_sig(x, h):
    f = np.array([sig(i,h)*2*cn(i)*np.exp(1j*2*i*np.pi*x/period) for i in range(0,h+1)])
    return f.sum()

def mod_sig(a, b, c, h): #Построение графика 
    xm = np.arange(a, b+c, c)
    y1 = np.array([f_okon_sig(x, h).real for x in xm])
    y2 = np.array([f_okon(x, h, priamougl_1).real for x in xm])
    plt.plot(xm, y1, 'r-', linewidth = 4) 
    plt.plot(xm, y2, 'b-', linewidth = 4, alpha = 0.6)
    plt.grid()
    plt.show() 
    print(analys(xm, y1))
    print(analys_l1_li(xm, y1))
    return None

#-------------------------------------------------------------------------------------------------------------

def lagrange(x, xx): #Вспомогательная функция для построения полинома Лагранжа
    n=len(x); m=len(xx)
    L=np.zeros((m,n))
    for k in range(n):
        x_k = np.concatenate((x[:k],x[k+1:]))
        L[:,k] = np.prod(np.matlib.repmat(xx.reshape(m,1),1,n-1) - np.matlib.repmat(x_k,m,1) , axis=1) / np.prod(x[k]-x_k)
    return L

def lagrange_interp(x,y,xx): # Функция для построения полинома Лагранжа на узлах (x,y) и точках xx
    L = lagrange(x,xx)
    return np.dot(L,y)

def ChebyL(n, a, b): # узлы Чебышева-Лоббато на отрезке [a;b]
    return ((a+b)/2-((a-b)/2)*np.cos(np.pi*np.arange(0.,n)/(n-1)))

def ChebyG(n): # узлы Чебышева-Гаусса на отрезке [-1;1]
    return np.cos(np.pi*(2*np.arange(0.,n)+1)/(2*n))

def runge_cheak(a, b, c):
    f = lambda x : 1./(25*x**2 + 1)
    x_test  = np.arange(a, b+c, c)   
    y_true = f(x_test)
    
    #x_train5 = np.linspace(a,b,5)
    #x_train15 = np.linspace(a,b,15)
    x_train25 = np.linspace(a,b,11)
    #y_train5 = f(x_train5)
    #y_train15 = f(x_train15)
    y_train25 = f(x_train25)

    #y_poly5 = lagrange_interp(x_train5,y_train5, x_test)
    #y_poly15 = lagrange_interp(x_train15,y_train15, x_test)
    y_poly25 = lagrange_interp(x_train25,y_train25, x_test)
    
    plt.plot(x_test, y_poly25, 'r-', linewidth = 4)
    plt.plot(x_test, y_true, 'b-', linewidth = 4)
    plt.grid() 
    plt.show()
    return None

def main(N, a, b, c): # Расчет характеристик для сравнения трех методов подавления явления Рунге
    f = lambda x : 1./(25*x**2 + 1)
    S = lambda x : -(a + b)/2 + ((a - b)/2)*np.cos(np.pi*(x-a)/(b-a))
    x_train = np.linspace(a,b,N)
    x_test  = np.arange(a, b+c, c)
    
    x_fake_train = S(x_train)
    x_fake_test = S(x_test)

    y_train = f(x_train)   
    y_true = f(x_test)  
    x_CL1  = ChebyL(N, a, b)
    y_CL1 = f(x_CL1)
    x_CG = ChebyG(N)
    y_CG = f(x_CG)
    y_ChebL = lagrange_interp(x_CL1, y_CL1, x_test)
    #y_ChebG = lagrange_interp(x_CG, y_CG, x_test)
    y_fake = lagrange_interp(x_fake_train, y_train, x_fake_test)
    y_poly = lagrange_interp(x_train,y_train, x_test)
    
    fig, ax = plt.subplots(1,1, figsize=(14,14))
    #plt.plot(x_test, y_poly,'g-',x_test,y_true,'b--', x_train, y_train,'k.', x_test, y_Cheb,'y-', x_CL, y_CL,'r.') 
    #plt.plot(x_test, y_fake, 'r-', x_test, y_poly,'b-', x_test, y_ChebL, 'k-', x_test, y_true,'y-' )
    plt.plot(x_test, y_true, 'b-', linewidth = 1)
    plt.plot(x_fake_test, y_true, 'r-', linewidth = 1)
    plt.plot(x_train, y_train, 'go', linewidth = 4)
    plt.plot(x_fake_train, y_train, 'ro')
    #plt.plot(x_CL1, y_CL1, 'go', linewidth = 70)
    #plt.plot(x_CG, y_CG, 'r^', linewidth = 70)
    
    ax.arrow(x_train[1] + x_train[1]/65, y_train[1], -(x_train[1]- x_fake_train[1])*0.80 , 0, width = 0.009, head_length = 0.01)
    ax.arrow(x_train[2] + x_train[2]/35, y_train[2], -(x_train[2]- x_fake_train[2])*0.80 , 0, width = 0.009, head_length = 0.01)
    ax.arrow(x_train[3] + x_train[3]/20, y_train[3], -(x_train[3]- x_fake_train[3])*0.68 , 0, width = 0.009, head_length = 0.01)
    ax.arrow(x_train[5] + x_train[5]/20, y_train[5], -(x_train[5]- x_fake_train[5])*0.68 , 0, width = 0.009, head_length = 0.01)
    ax.arrow(x_train[6] + x_train[6]/35, y_train[6], -(x_train[6]- x_fake_train[6])*0.80 , 0, width = 0.009, head_length = 0.01)
    ax.arrow(x_train[7] + x_train[7]/65, y_train[7], -(x_train[7]- x_fake_train[7])*0.80 , 0, width = 0.009, head_length = 0.01)
    plt.grid() 
    plt.show()
    res = np.array([])
    res = np.append(res, round(N))
    res = np.append(res, round(np.linalg.norm(y_true - y_poly,1), 6))
    res = np.append(res, round(np.linalg.norm(y_true - y_ChebL,1), 6))
    res = np.append(res, round(np.linalg.norm(y_true - y_fake,1), 6))
    res = np.append(res, round(np.linalg.norm(y_true - y_poly, np.inf), 6))
    res = np.append(res, round(np.linalg.norm(y_true - y_ChebL, np.inf), 6))
    res = np.append(res, round(np.linalg.norm(y_true - y_fake, np.inf), 6))
    return res

def main1(a, b, c): # Создание word файла с таблицой характеристик для разного количества узлов
    res1 = main(5, a, b, c)
    res2 = main(7, a, b, c)
    res3 = main(10, a, b, c)
    res4 = main(15, a, b, c)
    res5 = main(25, a, b, c)
    allres = np.array([res1, res2, res3, res4, res5])
    doc = docx.Document()
    table = doc.add_table(rows = 6, cols = 7)  
    table.style = 'Table Grid'
    for rows in range(1, 6):
        for col in range(7):
            cell = table.cell(rows, col)
            cell.text = str(allres[rows-1][col]) 
    doc.save('zxc.docx') 
    return None